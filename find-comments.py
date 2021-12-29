import docx
from docx import Document


doc = Document("/home/anmol/Downloads/MSA2.docx")
print(doc.paragraphs[0].text)

# loop to get all paragraphs

index=0
for para in doc.paragraphs:             
    index=index+1                 
    if(len(para.text)>0):
        print("paragraph",index,"is")
        print(para.text)

# position of comment
from docx import Document
from lxml import etree
import zipfile
ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
#Function to extract all the comments of document(Same as accepted answer)
#Returns a dictionary with comment id as key and comment string as value
def get_document_comments(docxFileName):
    comments_dict={}
    docxZip = zipfile.ZipFile(docxFileName)
    commentsXML = docxZip.read('word/comments.xml')
    et = etree.XML(commentsXML)
    comments = et.xpath('//w:comment',namespaces=ooXMLns)
    for c in comments:
        comment=c.xpath('string(.)',namespaces=ooXMLns)
        comment_id=c.xpath('@w:id',namespaces=ooXMLns)[0]
        author=(c.xpath('@w:author',namespaces=ooXMLns))
        date=(c.xpath('@w:date',namespaces=ooXMLns))
        comments_dict[comment_id]=[comment,author,date]
    return comments_dict
#Function to fetch all the comments in a paragraph
def paragraph_comments(paragraph,comments_dict):
    comments=[]
    for run in paragraph.runs:
        comment_reference=run._r.xpath("./w:commentReference")
        if comment_reference:
            comment_id=comment_reference[0].xpath('@w:id',namespaces=ooXMLns)[0]
            # author=(comment_reference[0].xpath('@w:author',namespaces=ooXMLns))
            # date=(comment_reference[0].xpath('@w:date',namespaces=ooXMLns))
            comment=comments_dict[comment_id]
            comments.append(comment)
    return comments
#Function to fetch all comments with their referenced paragraph
#This will return list like this [{'Paragraph text': [comment 1,comment 2]}]
def comments_with_reference_paragraph(docxFileName):
    document = Document(docxFileName)
    comments_dict=get_document_comments(docxFileName)
    comments_with_their_reference_paragraph=[]
    for paragraph in document.paragraphs:  
        if comments_dict: 
            comments=paragraph_comments(paragraph,comments_dict)  
            if comments:
                comments_with_their_reference_paragraph.append({paragraph.text: comments})
    return comments_with_their_reference_paragraph
if __name__=="__main__":
    document="/home/anmol/Downloads/MSA2.docx"  #filepath for the input document
    print(comments_with_reference_paragraph(document))
